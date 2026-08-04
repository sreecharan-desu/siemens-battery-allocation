#!/usr/bin/env python3
"""Generate Siemens Battery Allocation presentation (Word + PDF)."""

from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RlImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

BASE = Path(__file__).parent.parent
OUTPUTS = BASE / "outputs"
DOCS = Path(__file__).parent
OUTPUT_DOCX = DOCS / "Siemens_Battery_Allocation_Presentation.docx"
OUTPUT_PDF = DOCS / "Siemens_Battery_Allocation_Presentation.pdf"

ACCENT = RGBColor(0x1A, 0x36, 0x5D)
TEXT = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)

MAX_IMG_WIDTH = Inches(6.0)
MAX_IMG_HEIGHT = Inches(3.8)


def load_metrics() -> dict:
    path = OUTPUTS / "metrics_report.json"
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {}


def set_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    for level, size in [(1, 22), (2, 16), (3, 13)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.bold = True
        h.font.color.rgb = ACCENT if level == 1 else TEXT
        h.font.size = Pt(size)
        h.paragraph_format.space_before = Pt(12 if level > 1 else 18)
        h.paragraph_format.space_after = Pt(6)


def picture_size(path: Path) -> tuple:
    with Image.open(path) as im:
        px_w, px_h = im.size
    aspect = px_h / px_w
    w = MAX_IMG_WIDTH
    h = w * aspect
    if h > MAX_IMG_HEIGHT:
        h = MAX_IMG_HEIGHT
        w = h / aspect
    return w, h


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_cover_page(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Battery Health Assessment\nand Dynamic Allocation")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Siemens Energy · IMECE India 2026\nBrain Bolt Engineers Sprint")
    s.font.size = Pt(14)
    s.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(
        f"Prepared {date.today().strftime('%d %B %Y')}  ·  "
        "github.com/sreecharan-desu/siemens-battery-allocation"
    )
    m.font.size = Pt(10)
    m.font.color.rgb = MUTED

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    v = ver.add_run("battery-allocation v1.2.1")
    v.font.size = Pt(10)
    v.font.color.rgb = MUTED

    doc.add_page_break()


def add_executive_summary(doc: Document, metrics: dict) -> None:
    proposed = metrics.get("proposed", {})
    baseline = metrics.get("baseline", {})

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This document presents a production-grade system for classifying battery health, "
        "scoring pack suitability, and allocating batteries to light EV swap requests at a "
        "station with 200 packs and 50 concurrent vehicle demands. The proposed "
        "Priority-Suitability allocator matches the Highest-SoC-First baseline on vehicles "
        f"served ({proposed.get('vehicles_served', 36)}/50) while improving high/critical "
        f"priority coverage (+{proposed.get('high_critical_served_pct', 73.08) - baseline.get('high_critical_served_pct', 69.23):.2f} pp), "
        f"average state-of-health (+{proposed.get('avg_soh_allocated', 90.37) - baseline.get('avg_soh_allocated', 80.49):.2f}%), "
        "and suitability scores — with zero unsafe allocations and zero constraint violations."
    )

    doc.add_heading("Platform Overview", level=2)
    add_bullets(doc, [
        "Classifies 200 batteries into Safe, Degraded, or Unsafe/Quarantine (14 quarantined)",
        "Scores each pack 0–100 on five health parameters (SOH, SOC, resistance, imbalance, temperature)",
        "Allocates batteries by vehicle priority (Critical → High → Normal) and suitability",
        "Compares against Highest-SoC-First baseline on every run",
        "Exports CSV tables, JSON metrics, and five analytical charts",
        "CLI, REST API, Docker, and CI pipeline with 58 automated tests",
    ])

    doc.add_heading("Key Results", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Metric", "Proposed", "Baseline", "Delta"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True

    rows = [
        ("Vehicles served", proposed.get("vehicles_served"), baseline.get("vehicles_served"), "0"),
        ("High/Critical served %", proposed.get("high_critical_served_pct"), baseline.get("high_critical_served_pct"), "+3.85"),
        ("Avg SoH allocated", proposed.get("avg_soh_allocated"), baseline.get("avg_soh_allocated"), "+9.88"),
        ("Avg suitability score", proposed.get("avg_suitability_score"), baseline.get("avg_suitability_score"), "+7.74"),
        ("Unsafe allocations", proposed.get("unsafe_allocations"), baseline.get("unsafe_allocations"), "0"),
    ]
    for label, p, b, d in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(p)
        row[2].text = str(b)
        row[3].text = str(d)

    doc.add_paragraph()
    doc.add_page_break()


def add_figure_section(
    doc: Document,
    number: int,
    title: str,
    filename: str,
    caption: str,
    observations: list[str],
) -> None:
    doc.add_heading(f"{number}. {title}", level=2)

    img_path = OUTPUTS / filename
    if img_path.exists():
        w, h = picture_size(img_path)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img_path), width=w, height=h)

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(f"Figure {number}: {caption}")
        cap_run.italic = True
        cap_run.font.size = Pt(10)
        cap_run.font.color.rgb = MUTED
    else:
        doc.add_paragraph(f"[Chart not found: {filename}]")

    doc.add_heading("Key Observations", level=3)
    add_bullets(doc, observations)
    doc.add_paragraph()


def add_visualizations(doc: Document) -> None:
    doc.add_heading("Results & Visualizations", level=1)
    doc.add_paragraph(
        "The following figures were generated from the competition datasets "
        "(200 battery packs, 50 vehicle requests) using battery-allocation run. "
        "All charts are reproducible from the submitted code and input CSV files."
    )

    sections = [
        (
            1,
            "Battery Classification Distribution",
            "01_battery_classification.png",
            "Distribution of 200 battery packs across Safe, Degraded, and Unsafe/Quarantine categories.",
            [
                "128 batteries (64%) classified as Safe and Available for allocation",
                "58 batteries (29%) classified as Degraded but Usable under constraints",
                "14 batteries (7%) flagged Unsafe/Quarantine and excluded from all allocations",
                "Classification uses configurable thresholds in config/thresholds.yaml",
                "Station status REVIEW/QUARANTINE triggers immediate unsafe classification",
            ],
        ),
        (
            2,
            "Suitability Score Distribution",
            "02_suitability_score_distribution.png",
            "Histogram of suitability scores (0–100) grouped by battery category.",
            [
                "Safe packs cluster at higher suitability scores (avg ~75.6)",
                "Degraded packs show wider spread (avg ~61.8) reflecting mixed health signals",
                "Unsafe/quarantined packs still scored for audit but never allocated",
                "Score weights: SOH 30%, SOC 25%, resistance 20%, imbalance 15%, temperature 10%",
            ],
        ),
        (
            3,
            "Allocation by Priority (Proposed Method)",
            "03_allocation_by_priority_proposed.png",
            "Vehicles served vs unserved broken down by Critical, High, and Normal priority.",
            [
                "Critical: 2 of 4 served (50.0%) — limited by eligible pack availability",
                "High: 17 of 22 served (77.3%) — strongest coverage among priority tiers",
                "Normal: 17 of 24 served (70.8%)",
                "Priority ordering ensures Critical and High requests are processed first",
            ],
        ),
        (
            4,
            "Method Comparison — Proposed vs Baseline",
            "04_method_comparison.png",
            "Side-by-side comparison of Priority-Suitability vs Highest-SoC-First metrics.",
            [
                "Both methods serve 36 of 50 vehicles — same throughput",
                "Proposed method improves High/Critical served rate by 3.85 percentage points",
                "Average SoH of allocated packs rises from 80.5% to 90.4% (+9.88 pp)",
                "Average suitability score rises from 76.8 to 84.5 (+7.74 points)",
                "All 36 served vehicles receive different battery assignments under proposed method",
            ],
        ),
        (
            5,
            "Quarantine Battery Identification",
            "05_quarantine_batteries.png",
            "Visualization of 14 unsafe/quarantined batteries excluded from allocation.",
            [
                "Each quarantined pack identified by battery_id with triggering health parameters",
                "Full detail exported to outputs/quarantine_report.csv",
                "Zero unsafe batteries appear in proposed or baseline allocation outputs",
                "Constraint verifier runs automatically after every pipeline execution",
            ],
        ),
    ]

    for args in sections:
        add_figure_section(doc, *args)


def add_capability_matrix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("System Capability Matrix", level=1)
    doc.add_paragraph(
        "Capabilities delivered in the submission against typical competition and "
        "pilot-deployment requirements."
    )

    features = [
        ("Battery health classification (3 categories)", "Yes", "Delivered"),
        ("Suitability scoring 0–100", "Yes", "Delivered"),
        ("Priority-aware allocation", "Yes", "Delivered"),
        ("Highest-SoC-First baseline comparison", "Yes", "Delivered"),
        ("Constraint verification (no unsafe, no duplicates, SOC min)", "Yes", "Delivered"),
        ("CSV + Excel data ingestion", "Yes", "Delivered"),
        ("Dynamic file discovery (no hardcoded paths)", "Yes", "Delivered"),
        ("Interactive CLI", "Yes", "Delivered"),
        ("REST API (FastAPI)", "Yes", "Delivered"),
        ("Docker container", "Yes", "Delivered"),
        ("CI pipeline (lint, type-check, tests)", "Yes", "Delivered"),
        ("Onsite twist handler", "Yes", "Delivered"),
        ("Configurable thresholds (YAML)", "Yes", "Delivered"),
        ("Automated visualizations (5 charts)", "Yes", "Delivered"),
        ("Cross-platform setup (macOS/Linux/Windows)", "Yes", "Delivered"),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Capability", "Status", "Notes"]):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True

    for feat, status, notes in features:
        row = table.add_row().cells
        row[0].text = feat
        row[1].text = status
        row[2].text = notes

    doc.add_paragraph()


def add_method_section(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Method Summary", level=1)

    doc.add_heading("Classification Rules", level=2)
    add_bullets(doc, [
        "Unsafe/Quarantine: station REVIEW/QUARANTINE, or SOH/temp/resistance/imbalance beyond unsafe thresholds",
        "Degraded but Usable: not unsafe, but exceeds degraded thresholds on SOH, resistance, imbalance, or cycles",
        "Safe and Available: all parameters within safe limits",
    ])

    doc.add_heading("Proposed Allocation (Priority-Suitability)", level=2)
    add_bullets(doc, [
        "Sort requests by priority (Critical → High → Normal), then arrival time",
        "For each request, filter eligible batteries: not unsafe, SOC ≥ minimum, energy ≥ trip requirement",
        "Score candidates: suitability + energy margin + category bonus + SOC balance",
        "Assign highest-scoring unused battery",
    ])

    doc.add_heading("Architecture", level=2)
    doc.add_paragraph(
        "Layered Python package: CLI and REST API → Pipeline Runner → "
        "Classification / Scoring / Allocation / Twist Handler → Data Loader → YAML config. "
        "Deployment modes: batch CLI, API server, Docker container."
    )

    doc.add_heading("Quality Gates", level=2)
    add_bullets(doc, [
        "58 automated tests, ≥80% code coverage",
        "Ruff lint and Mypy strict mode in CI",
        "Pipeline smoke test on Python 3.11 and 3.12",
        "All constraint checks pass on competition datasets",
    ])


def add_recommendations(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Recommendations", level=1)

    doc.add_heading("Competition Presentation", level=2)
    add_bullets(doc, [
        "Lead with safety: 0 unsafe allocations, 14 batteries correctly quarantined",
        "Highlight method comparison chart — same serve count, better health outcomes",
        "Demonstrate live CLI: battery-allocation run --sample",
        "Show twist handler for onsite scenario adaptation",
    ])

    doc.add_heading("Pilot Deployment", level=2)
    add_bullets(doc, [
        "Deploy API via Docker at station edge; integrate with BMS telemetry feed",
        "Tune thresholds.yaml on local fleet data before production cutover",
        "Use upload workflow for operator-submitted CSV/Excel exports from station systems",
        "Enable JSON logging for centralized monitoring",
    ])

    doc.add_heading("Future Enhancements", level=2)
    add_bullets(doc, [
        "Real-time streaming allocation as requests arrive (event-driven API)",
        "ML-based SOH prediction layered on rule-based safety guardrails",
        "Operator dashboard for quarantine review and manual override audit trail",
        "Multi-station fleet balancing across swap network",
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Document —")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Internal / competition submission reference. Proprietary — Siemens Energy / IMECE India 2026.")
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = MUTED


def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "Title",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=22,
            textColor=colors.HexColor("#1A365D"),
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "subtitle": ParagraphStyle(
            "Subtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=12,
            textColor=colors.HexColor("#5A5A5A"),
            alignment=TA_CENTER,
            spaceAfter=6,
        ),
        "h1": ParagraphStyle(
            "H1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            textColor=colors.HexColor("#1A365D"),
            spaceBefore=14,
            spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "H2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            spaceBefore=10,
            spaceAfter=6,
        ),
        "h3": ParagraphStyle(
            "H3",
            parent=base["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=11,
            spaceBefore=8,
            spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "Body",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            spaceAfter=6,
        ),
        "bullet": ParagraphStyle(
            "Bullet",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=13,
            leftIndent=18,
            bulletIndent=8,
            spaceAfter=3,
        ),
        "caption": ParagraphStyle(
            "Caption",
            parent=base["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=9,
            textColor=colors.HexColor("#5A5A5A"),
            alignment=TA_CENTER,
            spaceAfter=10,
        ),
    }


def _pdf_image(path: Path, max_w: float = 6.2 * inch) -> RlImage:
    with Image.open(path) as im:
        px_w, px_h = im.size
    aspect = px_h / px_w
    w = max_w
    h = w * aspect
    max_h = 3.6 * inch
    if h > max_h:
        h = max_h
        w = h / aspect
    return RlImage(str(path), width=w, height=h)


def _pdf_bullets(styles: dict[str, ParagraphStyle], items: list[str]) -> list:
    return [Paragraph(f"• {item}", styles["bullet"]) for item in items]


def _figure_sections() -> list[tuple]:
    return [
        (
            1,
            "Battery Classification Distribution",
            "01_battery_classification.png",
            "Distribution of 200 battery packs across Safe, Degraded, and Unsafe/Quarantine categories.",
            [
                "128 batteries (64%) classified as Safe and Available for allocation",
                "58 batteries (29%) classified as Degraded but Usable under constraints",
                "14 batteries (7%) flagged Unsafe/Quarantine and excluded from all allocations",
                "Classification uses configurable thresholds in config/thresholds.yaml",
                "Station status REVIEW/QUARANTINE triggers immediate unsafe classification",
            ],
        ),
        (
            2,
            "Suitability Score Distribution",
            "02_suitability_score_distribution.png",
            "Histogram of suitability scores (0–100) grouped by battery category.",
            [
                "Safe packs cluster at higher suitability scores (avg ~75.6)",
                "Degraded packs show wider spread (avg ~61.8) reflecting mixed health signals",
                "Unsafe/quarantined packs still scored for audit but never allocated",
                "Score weights: SOH 30%, SOC 25%, resistance 20%, imbalance 15%, temperature 10%",
            ],
        ),
        (
            3,
            "Allocation by Priority (Proposed Method)",
            "03_allocation_by_priority_proposed.png",
            "Vehicles served vs unserved broken down by Critical, High, and Normal priority.",
            [
                "Critical: 2 of 4 served (50.0%) — limited by eligible pack availability",
                "High: 17 of 22 served (77.3%) — strongest coverage among priority tiers",
                "Normal: 17 of 24 served (70.8%)",
                "Priority ordering ensures Critical and High requests are processed first",
            ],
        ),
        (
            4,
            "Method Comparison — Proposed vs Baseline",
            "04_method_comparison.png",
            "Side-by-side comparison of Priority-Suitability vs Highest-SoC-First metrics.",
            [
                "Both methods serve 36 of 50 vehicles — same throughput",
                "Proposed method improves High/Critical served rate by 3.85 percentage points",
                "Average SoH of allocated packs rises from 80.5% to 90.4% (+9.88 pp)",
                "Average suitability score rises from 76.8 to 84.5 (+7.74 points)",
                "All 36 served vehicles receive different battery assignments under proposed method",
            ],
        ),
        (
            5,
            "Quarantine Battery Identification",
            "05_quarantine_batteries.png",
            "Visualization of 14 unsafe/quarantined batteries excluded from allocation.",
            [
                "Each quarantined pack identified by battery_id with triggering health parameters",
                "Full detail exported to outputs/quarantine_report.csv",
                "Zero unsafe batteries appear in proposed or baseline allocation outputs",
                "Constraint verifier runs automatically after every pipeline execution",
            ],
        ),
    ]


def _capability_rows() -> list[tuple[str, str, str]]:
    return [
        ("Battery health classification (3 categories)", "Yes", "Delivered"),
        ("Suitability scoring 0–100", "Yes", "Delivered"),
        ("Priority-aware allocation", "Yes", "Delivered"),
        ("Highest-SoC-First baseline comparison", "Yes", "Delivered"),
        ("Constraint verification (no unsafe, no duplicates, SOC min)", "Yes", "Delivered"),
        ("CSV + Excel data ingestion", "Yes", "Delivered"),
        ("Dynamic file discovery (no hardcoded paths)", "Yes", "Delivered"),
        ("Interactive CLI", "Yes", "Delivered"),
        ("REST API (FastAPI)", "Yes", "Delivered"),
        ("Docker container", "Yes", "Delivered"),
        ("CI pipeline (lint, type-check, tests)", "Yes", "Delivered"),
        ("Onsite twist handler", "Yes", "Delivered"),
        ("Configurable thresholds (YAML)", "Yes", "Delivered"),
        ("Automated visualizations (5 charts)", "Yes", "Delivered"),
        ("Cross-platform setup (macOS/Linux/Windows)", "Yes", "Delivered"),
    ]


def generate_pdf(metrics: dict) -> None:
    proposed = metrics.get("proposed", {})
    baseline = metrics.get("baseline", {})
    styles = _pdf_styles()
    story: list = []

    story.append(Spacer(1, 1.8 * inch))
    story.append(Paragraph("Battery Health Assessment<br/>and Dynamic Allocation", styles["title"]))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Siemens Energy · IMECE India 2026<br/>Brain Bolt Engineers Sprint", styles["subtitle"]))
    story.append(Paragraph(
        f"Prepared {date.today().strftime('%d %B %Y')} · github.com/sreecharan-desu/siemens-battery-allocation",
        styles["subtitle"],
    ))
    story.append(Paragraph("battery-allocation v1.2.1", styles["subtitle"]))
    story.append(PageBreak())

    story.append(Paragraph("Executive Summary", styles["h1"]))
    story.append(Paragraph(
        "This document presents a production-grade system for classifying battery health, "
        "scoring pack suitability, and allocating batteries to light EV swap requests at a "
        "station with 200 packs and 50 concurrent vehicle demands. The proposed "
        "Priority-Suitability allocator matches the Highest-SoC-First baseline on vehicles "
        f"served ({proposed.get('vehicles_served', 36)}/50) while improving high/critical "
        f"priority coverage (+{proposed.get('high_critical_served_pct', 73.08) - baseline.get('high_critical_served_pct', 69.23):.2f} pp), "
        f"average state-of-health (+{proposed.get('avg_soh_allocated', 90.37) - baseline.get('avg_soh_allocated', 80.49):.2f}%), "
        "and suitability scores — with zero unsafe allocations and zero constraint violations.",
        styles["body"],
    ))

    story.append(Paragraph("Platform Overview", styles["h2"]))
    story.extend(_pdf_bullets(styles, [
        "Classifies 200 batteries into Safe, Degraded, or Unsafe/Quarantine (14 quarantined)",
        "Scores each pack 0–100 on five health parameters (SOH, SOC, resistance, imbalance, temperature)",
        "Allocates batteries by vehicle priority (Critical → High → Normal) and suitability",
        "Compares against Highest-SoC-First baseline on every run",
        "Exports CSV tables, JSON metrics, and five analytical charts",
        "CLI, REST API, Docker, and CI pipeline with 58 automated tests",
    ]))

    story.append(Paragraph("Key Results", styles["h2"]))
    table_data = [
        ["Metric", "Proposed", "Baseline", "Delta"],
        ["Vehicles served", str(proposed.get("vehicles_served")), str(baseline.get("vehicles_served")), "0"],
        ["High/Critical served %", str(proposed.get("high_critical_served_pct")), str(baseline.get("high_critical_served_pct")), "+3.85"],
        ["Avg SoH allocated", str(proposed.get("avg_soh_allocated")), str(baseline.get("avg_soh_allocated")), "+9.88"],
        ["Avg suitability score", str(proposed.get("avg_suitability_score")), str(baseline.get("avg_suitability_score")), "+7.74"],
        ["Unsafe allocations", str(proposed.get("unsafe_allocations")), str(baseline.get("unsafe_allocations")), "0"],
    ]
    t = Table(table_data, colWidths=[2.4 * inch, 1.1 * inch, 1.1 * inch, 0.9 * inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF4")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(t)
    story.append(PageBreak())

    story.append(Paragraph("Results &amp; Visualizations", styles["h1"]))
    story.append(Paragraph(
        "The following figures were generated from the competition datasets "
        "(200 battery packs, 50 vehicle requests) using battery-allocation run. "
        "All charts are reproducible from the submitted code and input CSV files.",
        styles["body"],
    ))

    for number, title, filename, caption, observations in _figure_sections():
        story.append(Paragraph(f"{number}. {title}", styles["h2"]))
        img_path = OUTPUTS / filename
        if img_path.exists():
            story.append(_pdf_image(img_path))
            story.append(Paragraph(f"Figure {number}: {caption}", styles["caption"]))
        else:
            story.append(Paragraph(f"[Chart not found: {filename}]", styles["body"]))
        story.append(Paragraph("Key Observations", styles["h3"]))
        story.extend(_pdf_bullets(styles, observations))
        story.append(Spacer(1, 0.15 * inch))

    story.append(PageBreak())
    story.append(Paragraph("System Capability Matrix", styles["h1"]))
    story.append(Paragraph(
        "Capabilities delivered in the submission against typical competition and pilot-deployment requirements.",
        styles["body"],
    ))
    cap_data = [["Capability", "Status", "Notes"], *_capability_rows()]
    cap_table = Table(cap_data, colWidths=[3.2 * inch, 0.7 * inch, 1.2 * inch])
    cap_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF4")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(cap_table)
    story.append(PageBreak())

    story.append(Paragraph("Method Summary", styles["h1"]))
    story.append(Paragraph("Classification Rules", styles["h2"]))
    story.extend(_pdf_bullets(styles, [
        "Unsafe/Quarantine: station REVIEW/QUARANTINE, or SOH/temp/resistance/imbalance beyond unsafe thresholds",
        "Degraded but Usable: not unsafe, but exceeds degraded thresholds on SOH, resistance, imbalance, or cycles",
        "Safe and Available: all parameters within safe limits",
    ]))
    story.append(Paragraph("Proposed Allocation (Priority-Suitability)", styles["h2"]))
    story.extend(_pdf_bullets(styles, [
        "Sort requests by priority (Critical → High → Normal), then arrival time",
        "For each request, filter eligible batteries: not unsafe, SOC ≥ minimum, energy ≥ trip requirement",
        "Score candidates: suitability + energy margin + category bonus + SOC balance",
        "Assign highest-scoring unused battery",
    ]))
    story.append(Paragraph("Architecture", styles["h2"]))
    story.append(Paragraph(
        "Layered Python package: CLI and REST API → Pipeline Runner → "
        "Classification / Scoring / Allocation / Twist Handler → Data Loader → YAML config. "
        "Deployment modes: batch CLI, API server, Docker container.",
        styles["body"],
    ))
    story.append(Paragraph("Quality Gates", styles["h2"]))
    story.extend(_pdf_bullets(styles, [
        "58 automated tests, ≥80% code coverage",
        "Ruff lint and Mypy strict mode in CI",
        "Pipeline smoke test on Python 3.11 and 3.12",
        "All constraint checks pass on competition datasets",
    ]))
    story.append(PageBreak())

    story.append(Paragraph("Recommendations", styles["h1"]))
    story.append(Paragraph("Competition Presentation", styles["h2"]))
    story.extend(_pdf_bullets(styles, [
        "Lead with safety: 0 unsafe allocations, 14 batteries correctly quarantined",
        "Highlight method comparison chart — same serve count, better health outcomes",
        "Demonstrate live CLI: battery-allocation run --sample",
        "Show twist handler for onsite scenario adaptation",
    ]))
    story.append(Paragraph("Pilot Deployment", styles["h2"]))
    story.extend(_pdf_bullets(styles, [
        "Deploy API via Docker at station edge; integrate with BMS telemetry feed",
        "Tune thresholds.yaml on local fleet data before production cutover",
        "Use upload workflow for operator-submitted CSV/Excel exports from station systems",
        "Enable JSON logging for centralized monitoring",
    ]))
    story.append(Paragraph("Future Enhancements", styles["h2"]))
    story.extend(_pdf_bullets(styles, [
        "Real-time streaming allocation as requests arrive (event-driven API)",
        "ML-based SOH prediction layered on rule-based safety guardrails",
        "Operator dashboard for quarantine review and manual override audit trail",
        "Multi-station fleet balancing across swap network",
    ]))
    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph("— End of Document —", styles["caption"]))
    story.append(Paragraph(
        "Internal / competition submission reference. Proprietary — Siemens Energy / IMECE India 2026.",
        styles["caption"],
    ))

    doc = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    doc.build(story)
    print(f"Saved: {OUTPUT_PDF}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate presentation documents")
    parser.add_argument("--docx-only", action="store_true", help="Generate Word document only")
    parser.add_argument("--pdf-only", action="store_true", help="Generate PDF only")
    args = parser.parse_args()

    metrics = load_metrics()
    if not args.pdf_only:
        doc = Document()
        set_document_styles(doc)
        add_cover_page(doc)
        add_executive_summary(doc, metrics)
        add_visualizations(doc)
        add_capability_matrix(doc)
        add_method_section(doc)
        add_recommendations(doc)
        doc.save(OUTPUT_DOCX)
        print(f"Saved: {OUTPUT_DOCX}")
    if not args.docx_only:
        generate_pdf(metrics)


if __name__ == "__main__":
    main()
